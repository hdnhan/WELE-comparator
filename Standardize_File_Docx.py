import docx
import unidecode
import string 
import re

def decontracted(phrase):
    # specific
    phrase = re.sub(r"won\'t", "will not", phrase)
    phrase = re.sub(r"can\'t", "can not", phrase)

    # general
    phrase = re.sub(r"n\'t", " not", phrase)
    phrase = re.sub(r"\'re", " are", phrase)
    phrase = re.sub(r"\'s", " is", phrase)
    phrase = re.sub(r"\'d", " would", phrase)
    phrase = re.sub(r"\'ll", " will", phrase)
    phrase = re.sub(r"\'t", " not", phrase)
    phrase = re.sub(r"\'ve", " have", phrase)
    phrase = re.sub(r"\'m", " am", phrase)
    return phrase

def Standardize_File_Docx(file_docx_name):
    #file = open(file_docx_name + '.docx', 'rb')
    file = open(file_docx_name, 'rb')
    content = docx.Document(file)
    standard_content = ''

    # transfer the whole document (many paragraphs) to only one paragraph
    for para in content.paragraphs:
        standard_content += para.text + ' '       
    
    # English and lower words
    standard_content = unidecode.unidecode(standard_content.lower().replace('\n', ' '))

    # no_space: replace letters with no space
    # one_space: replace letter with one space

    special_word = 'uh ah oh wow'

##    for elem in no_space.split():
##        standard_content = standard_content.replace(elem, '')
##
##    one_space = one_space.split() + ['    ', '   ', '  ']
##    for elem in one_space:
##        standard_content = standard_content.replace(elem, ' ')

    standard_content = decontracted(standard_content)

    standard_content = standard_content.translate(None,string.punctuation)


    
    for elem in special_word.split():
        # to ensure that dont remove substring of a word
        elem = ' ' + elem + ' '
        standard_content = standard_content.replace(elem, ' ')

    # to ensure all mistakes that are inside the paragraph
    extra = ' i love you '
    standard_content = extra + standard_content + extra
    return standard_content


def Save_File_Docx(content, file_name):
    doc = docx.Document()
    doc.add_paragraph(content)
    doc.save(file_name + ' Standard.docx')
