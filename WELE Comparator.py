import docx
import unidecode
import diff_match_patch as dmp_module
dmp = dmp_module.diff_match_patch()

# Desired results
Remove = []  # from orginial file (WELE file, answer)
Insert = []  # from compared file (Member's file)


def Standardize_File_Docx(file_docx_name):
    file = open(file_docx_name + '.docx', 'rb')
    content = docx.Document(file)
    standard_content = ''
    for para in content.paragraphs:
        standard_content += para.text + ' '
    standard_content = standard_content\
        .replace("'", '').replace('"', '')\
        .replace("‘", '').replace("’", '')\
        .replace('“', '').replace('”', '')\
        .replace('(', '').replace(')', '')\
        .replace(':', '')\
        .replace('.', '')\
        .replace('_', '')\
        .replace(',', '')\
        .replace('!', '')\
        .replace('?', '')\
        .replace('-', ' ')\
        .replace('–', ' ')\
        .replace('  ', ' ').replace('  ', ' ')\
        .replace('  ', ' ').replace('  ', ' ')\
        .replace('…', '')\
        .lower()
    return standard_content


esl_index = 10084
member_name = 'Ho Duc Nhan'
orginial_file_name = 'ESL' + ' ' + str(esl_index) + ' ' + 'WELE'
compared_file_name = 'ESL' + ' ' + str(esl_index) + ' ' + member_name


orginial_content = unidecode.unidecode(Standardize_File_Docx(orginial_file_name))
compared_content = unidecode.unidecode(Standardize_File_Docx(compared_file_name))

##################################################################

# save standard file
doc = docx.Document()
doc.add_paragraph(orginial_content)
doc.save(orginial_file_name + ' Standard.docx')
doc = docx.Document()
doc.add_paragraph(compared_content)
doc.save(compared_file_name + ' Standard.docx')

####################################################################
# to ensure all mistakes that are inside the content
extra_segment = ' i love you '
orginial_content = extra_segment + orginial_content + extra_segment
compared_content = extra_segment + compared_content + extra_segment
#################################################################

org_sp = orginial_content.split()
com_sp = compared_content.split()

sp_org = 'QWERTY '
sp_com = 'QWERTY '
for word in org_sp:
    sp_org += word + ' QWERTY '
for word in com_sp:
    sp_com += word + ' QWERTY '

#orginial_content = sp_org
#compared_content = sp_com

############################################################
#orginial_content = 'i love you babe bebe babe'
#compared_content = 'i love you bebe babe bebe'
extra_segment = ' i love you '
orginial_content = extra_segment + orginial_content + extra_segment
compared_content = extra_segment + compared_content + extra_segment
############################################################

patches = dmp.patch_make(orginial_content, compared_content)

###Start_Remove = []
###Count_Remove = []
Start_Insert = []
Count_Insert = []


for blocks in patches:
    lines = str(blocks).splitlines()
    index = lines[0].replace('@', '').replace(',', ' ')\
                    .replace('-', '').replace('+', '').split()
    ###Start_Remove.append(int(index[0]) - 1)
    # Count_Remove.append(int(index[1]))
    Start_Insert.append(int(index[2]) - 1)
    Count_Insert.append(int(index[3]))

###cnt = 0
# for it in range(len(Start_Remove) - 1):
###    cnt += Count_Remove[it] - Count_Insert[it]
###    Start_Remove[it + 1] += cnt


def Standardize_Block(block):
    x = 2
    while x < len(block) - 1:
        li_p = block[x - 1]
        li_c = block[x]
        li_n = block[x + 1]

        # delete this (remove or insert only one space)
        if li_c[0] != ' ' and len(li_c[1:].split()) == 0:
            block.pop(x)
            x -= 1
        # standardize the block
        # avoid a word seperating into two words
        # (note: more than 2: not resolved yet!!!)
        if(x + 2 < len(block) and li_p[0] != ' ' and li_p[0] == li_n[0] and li_c[0] == ' ' and block[x + 2][0] != ' '):
            # print('ok===>>>3')
            block[x - 1] += block[x][1:] + block[x + 1][1:]
            block[x + 2] = block[x + 2][0] + block[x][1:] + block[x + 2][1:]
            #print(block[x - 1])
            #print(block[x + 2])
            block.pop(x)
            block.pop(x)
            # print(block)
            x -= 2
        elif(li_p[0] != ' ' and li_c[0] != ' ' and li_n[0] != ' ' and li_p[-1] != ' '):
            # print('ok===>>>1')
            block[x - 1] += li_n[1:]
            block.pop(x + 1)
            # print(block)
            x -= 2
        elif(li_p[0] != ' ' and li_c[0] != ' ' and li_n[0] != ' ' and li_p[-1] == ' ' and x + 2 < len(block) and block[x + 2][1] == ' '):
            block[x - 1] += ' ' + block[x + 1][1:]
            block.pop(x + 1)
            x -= 1
            # print('ok===>>>')
            # print(block)
        elif(li_p[0] != ' ' and li_c[1] != ' ' and li_n[0] != ' ' and len(li_c.split()) == 1):
            # print('ok===>>>2')
            # print(li_c[1:])
            block[x - 1] += li_c[1:]
            block[x + 1] = li_n[0] + li_c[1:] + li_n[1:]
            #print(block[x - 1])
            #print(block[x + 1])
            block.pop(x)
            # print(block)
            x -= 1
        x += 1
    return block


# pos: position
# dire: direction
#   -1 => left
#    1 => right

def Get_Letters(pos, dire):
    res = ''
    i = 1
    while (compared_content[pos] != ' ' and compared_content[pos + i * dire] != ' '):
        if dire < 0:
            res = compared_content[pos + i * dire] + res
        else:
            res += compared_content[pos + i * dire]
        i += 1
    return res


# inner paragraphs
    # property: each block always has patches ' ' at both ends
    #           (meaning lines[begin or end].at[0] = ' ')

    # we have 4 patches' cases:
    # 1. '-'
    # 2. '+'
    # 3. '-' and '+'
    # 4. '+' and '-'

# fir: - or +
# sec: + or -


def detect(lines, it, fir, sec, pre_lef):
    line_prev = lines[it - 1]  # the privious line
    line_curr = lines[it]      # the current line
    line_next = lines[it + 1]  # the next line
    lp = line_prev.split()
    ln = line_next.split()

    if(line_curr[0] == fir and
        any([(line_curr[1] != ' '),
             (line_curr[1] == ' ' and len(line_curr) > 2)])):
        # => ignore the removing space (' ')

        # removing or inserting words in each block
        if((it == 2) and (len(lp) == 1) and line_prev[-1] != ' '):
            fir_words = pre_lef
            sec_words = pre_lef
        else:
            fir_words = ''
            sec_words = ''
        # case: '-+' or '+-'
        if (line_next[0] == sec):
            line_next_next = lines[it + 2]
            lnn = line_next_next.split()
            # line_prev.at(end) != space, meaning it's a part of a word
            if(line_prev[-1] != ' '):
                # add the final string of line_prev
                fir_words += lp[-1]
                sec_words += lp[-1]
            # add '-' to fir_words and '+' to sec_words
            fir_words += line_curr[1:]
            sec_words += line_next[1:]
            # a part of a word
            if(line_next_next[1] != ' '):
                if (line_curr[-1] != ' '):
                    fir_words += lnn[0]
                if (line_next[-1] != ' '):
                    sec_words += lnn[0]
        # case: '-'
        else:
            if(line_prev[-1] != ' '):
                fir_words += lp[-1]
                sec_words += lp[-1]
            if (fir == '-'):
                fir_words += line_curr[1:]
            else:
                sec_words += line_curr[1:]
            if(line_next[1] != ' ' and line_curr[-1] != ' '):
                fir_words += ln[0]
                sec_words += ln[0]

        Remove.append(fir_words)
        Insert.append(sec_words)


def inner(lines, pre_lef, nex_rig):
    it = 2
    while it < len(lines) - 1:  # ignoring the first and end lines
        # standardize lines here!!!

        line_curr = lines[it]
        if (line_curr[0] == '-'):
            detect(lines, it, '-', '+', pre_lef)
        elif(line_curr[0] == '+'):
            detect(lines, it, '+', '-', pre_lef)
        line_next = lines[it + 1]
        if(any([line_curr[0] == '-', line_curr[0] == '+']) and
                any([line_next[0] == '-', line_next[0] == '+'])):
            it += 1  # jump to the line behind the next line
        it += 1

    line_end = lines[-1]
    line_near_end = lines[-2]
    if(len(line_end.split()) == 1 and line_near_end[-1] != ' ' and line_end[1] != ' '):
        Remove[-1] += nex_rig
        Insert[-1] += nex_rig


for cnt in range(len(patches)):
    pre_lef = Get_Letters(Start_Insert[cnt], -1)
    nex_rig = Get_Letters(Start_Insert[cnt] + Count_Insert[cnt] - 1, 1)

    block = str(patches[cnt]).splitlines()

    # print(block)
    block = Standardize_Block(block)
    block = Standardize_Block(block)
    # print(block)

    inner(block, pre_lef, nex_rig)


##########################################################################
# RESULT

#Remove = [word for pair in Remove for word in pair.split() if  word.islower()]
#Insert = [word for pair in Insert for word in pair.split() if  word.islower()]
print('=' * 70)
print('|| ' + 'Remove     '.rjust(30) +
      ' || ' + '     Insert'.ljust(30) + ' ||')
print('=' * 70)
for it in range(len(Remove)):
    print('|| ' + Remove[it].rjust(30) + ' || ' + Insert[it].ljust(30) + ' ||')
print('=' * 70)
